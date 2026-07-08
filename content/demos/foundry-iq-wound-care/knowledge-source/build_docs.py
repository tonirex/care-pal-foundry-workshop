"""Build the 3 wound-care knowledge-base Word documents for the Foundry IQ demo.

Content is faithfully curated from the public source articles (nav/CTA cruft removed,
one garbled source sentence lightly repaired, one duplicate/mislabelled heading corrected
to match its content). Each .docx embeds its source URL so a grounded agent can cite it.

This script IS the editable source of the .docx files. A content owner / clinician can adjust
the wording below and regenerate before go-live:

    pip install python-docx
    python build_docs.py

The .docx are written next to this script (the `knowledge-source/` folder).
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
RETRIEVED = date.today().isoformat()

# Block grammar: (kind, text)
#   h1, h2  -> headings; p -> paragraph; b -> bullet; b2 -> nested bullet
#   quote   -> block quote (italic); note -> small grey note

HEALTHHUB = {
    "filename": "healthhub-wound-care.docx",
    "title": "Wound Care",
    "source_title": "HealthHub — Wound Care",
    "source_url": "https://www.healthhub.sg/health-conditions/wound-care",
    "publisher": "HealthHub (Ministry of Health, Singapore). "
                 "Contributed by the Pharmaceutical Society of Singapore. Last reviewed 23 January 2025.",
    "blocks": [
        ("h1", "What is a wound?"),
        ("p", "Wounds may refer to any type of damage or breakage of the tissues of your body. "
              "A wound may break the surface of the skin, but it can also occur without breaking the skin."),

        ("h1", "What are the possible causes of this condition?"),
        ("p", "Most wounds are caused by direct physical injury to your body. This may include accidents "
              "like burns, cuts, falling down, or intentional wounds like surgery and piercings. Some "
              "diseases may cause you to develop wounds more easily."),
        ("p", "Skin diseases like eczema and psoriasis make the skin more fragile and prone to breaking down."),
        ("p", "Diseases like diabetes can lead to poor blood circulation and numbness or loss of sensation, "
              "which makes it more difficult for you to notice if you have hurt yourself."),

        ("h1", "What are the possible types of wounds?"),
        ("p", "Wounds can include, but are not limited to, the following:"),
        ("b", "Abrasions"), ("b", "Bites"), ("b", "Bruises"), ("b", "Burns"), ("b", "Cuts"),
        ("b", "Incisions"), ("b", "Lacerations (deep cut or tear in skin)"), ("b", "Punctures"),
        ("b", "Skin tears"), ("b", "Sprains"), ("b", "Stings"), ("b", "Ulcers"),

        ("h1", "What can I do to treat wounds?"),
        ("p", "Some wounds require treatment by a healthcare professional. Other wounds are minor in nature "
              "and can be self-treated. Minor wounds can be treated in the following ways. "
              "(For burns, and for insect bites and stings, please refer to the separate articles.)"),
        ("h2", "For bruises and sprains"),
        ("b", "Apply a cold compress to the area for 15 to 20 minutes each time, three to four times a day "
              "for the first 48 to 72 hours."),
        ("b", "If possible, raise the injured area above the level of the heart."),
        ("b", "For sprains and swelling with a bruise, put an elastic bandage around it to compress the area. "
              "Take care not to wrap the bandage too tightly to avoid cutting off blood circulation."),
        ("b", "Take a painkiller such as paracetamol, ibuprofen or naproxen if necessary to manage pain."),
        ("b", "Rest the affected area as much as possible."),
        ("b", "For bruises, you may consider applying remedies containing Heparinoid to aid in healing."),
        ("h2", "For small cuts, incisions and abrasions"),
        ("b", "Cleanse the area with clean water or normal saline to remove any foreign particles."),
        ("b", "Apply an appropriate dressing and firm pressure to the wound area to stop the bleeding "
              "if necessary."),
        ("b", "To reduce the risk of infection, apply an antiseptic cream or solution to the affected area. "
              "DO NOT APPLY ALCOHOL ON BROKEN SKIN as it can lead to unnecessary pain and delay wound healing."),
        ("b", "Examples of possible antiseptics include:"),
        ("b2", "Chlorhexidine solution"),
        ("b2", "Octenidine"),
        ("b2", "Chlorhexidine cream"),
        ("b2", "Chloroxylenol"),
        ("b2", "Povidone Iodine"),
        ("b", "Cover the wound with a suitable dressing to reduce pain and prevent bacteria from entering "
              "the wound site."),
        ("b", "The dressing should be able to keep the wound site moist, but not overly wet, in order to give "
              "the best possible chance for the wound to heal without scarring."),
        ("b", "Do not try to dry the wound with baby powder. Powder can cause irritation and introduce "
              "bacteria into the wound."),
        ("b", "Do not apply medicated oil into the wound. Medicated oil will further damage the tissues in "
              "the injured area and delay wound healing."),

        ("h1", "When do I need to see a doctor?"),
        ("p", "Although minor wounds may be treated without a doctor's consultation, there are times where "
              "the condition might be more serious."),
        ("p", "If the wound is large or complicated. This includes but is not limited to:"),
        ("b", "Lacerations or loss of limb(s)"),
        ("b", "Wounds over a large area (more than 10%) of the body"),
        ("b", "Gaping wounds"),
        ("b", "Wounds where muscle, bone or fatty tissue can be seen"),
        ("b", "Wounds which have blood spurting from them"),
        ("b", "Deep puncture wounds"),
        ("b", "Animal bites"),
        ("b", "Burns which are more severe than a second degree burn / partial thickness burn"),
        ("b", "Wounds that cause severe pain, numbness or an inability to move the body part(s) affected "
              "by the wound"),
        ("b", "Wounds with visible dirt or other foreign materials (e.g. splinter) stuck inside even after "
              "washing the wound"),
        ("p", "Even if your wound is minor, you are still advised to see a healthcare professional if you "
              "have the following conditions:"),
        ("b", "A history of diabetes"),
        ("b", "Diseases that weaken your immune system (e.g. cancer, HIV)"),
        ("b", "Have not had a tetanus vaccination in the past 5 years"),
        ("p", "Even if your wound was initially minor, do also seek a doctor when the following occurs:"),
        ("b", "The wound shows no signs of improvement after more than a week has passed"),
        ("b", "Signs of infection are observed, such as: fever, swollen lymph nodes, pus, increasing or "
              "spreading redness around the wound area, swelling around the wound area, or warmth and pain "
              "around the wound area which is worsening"),

        ("h1", "What else can I do to manage this condition?"),
        ("p", "In addition to the steps mentioned above, here are a few more tips which can help to ensure "
              "that your wounds heal quickly and without scarring:"),
        ("b", "Keep the wound clean and covered to prevent it from becoming infected."),
        ("b", "Consume adequate nutrition to ensure that your body has sufficient resources to heal itself."),
        ("b", "Quit smoking."),
        ("b", "Refrain from over-exerting or putting pressure on the wounded area."),
        ("b", "If you have diabetes, ensure that your blood sugar is well controlled."),

        ("h1", "Disclaimer"),
        ("p", "This article is jointly developed by members of the National Medication Information "
              "workgroup. The content is solely for informational purposes only and is not intended as a "
              "substitute for the advice provided by your physician, pharmacist or other healthcare "
              "professional. Always speak with your physician, pharmacist or other healthcare professional "
              "before taking any medication or supplement, or adopting any treatment for a health problem."),
    ],
}

SLH_TREATMENT = {
    "filename": "slh-wound-care-treatment.docx",
    "title": "Wound Care Treatment in Singapore: How to Heal Faster and Avoid Complications",
    "source_title": "St Luke's Hospital — Wound Care Treatment in Singapore",
    "source_url": "https://www.slh.org.sg/wound-care-treatment-in-singapore-how-to-heal-faster-and-avoid-complications/",
    "publisher": "St Luke's Hospital, Singapore (St Luke's Wound Clinic).",
    "blocks": [
        ("p", "If you're dealing with a slow-healing wound, pressure injury, or diabetic ulcer, proper care "
              "is essential to prevent infection and ensure a smooth recovery. At St Luke's Wound Clinic in "
              "Singapore, we specialise in managing both acute and chronic wounds with a team of dedicated "
              "experts."),

        ("h1", "Why proper wound care matters"),
        ("p", "Ignoring even a minor wound can lead to serious problems such as infection, delayed healing, "
              "or even sepsis. With the right wound care, you can:"),
        ("b", "Heal faster with less pain"),
        ("b", "Reduce risk of infection"),
        ("b", "Minimise scarring"),
        ("b", "Avoid costly hospital visits or complications"),

        ("h1", "How to treat a wound at home (basic first aid)"),
        ("b", "Wash your hands to avoid bacteria."),
        ("b", "Clean the wound gently with saline or mild soap and water."),
        ("b", "Apply antibiotic ointment and cover with a sterile dressing."),
        ("b", "Change dressings daily or when they become wet or dirty."),
        ("b", "Watch for infection — redness, swelling, pus, and pain are warning signs."),

        ("h1", "When to see a wound care specialist"),
        ("p", "Seek professional wound care if any of the following apply:"),
        ("b", "Your wound is deep, large, or not healing after a few days."),
        ("b", "There are signs of infection (fever, pus, swelling)."),
        ("b", "You have a chronic condition like diabetes."),
        ("b", "You're caring for someone bedridden or immobile who has developed wounds."),
        ("b", "The wound was caused by rusty metal, an animal bite, or surgery."),

        ("h1", "Expert insight"),
        ("quote", "Taking good care of your wound means keeping it protected and watching for any changes "
                  "that might require a wound clinician's attention. Healing also comes from within — eat "
                  "nutritious foods, stay well hydrated, and avoid smoking to give your body the best chance "
                  "to repair itself. Remember: \u201cCare for your wound, care for yourself.\u201d"),
        ("note", "— Kavitha D/O Sanmugam, Senior Wound Clinician, St Luke's Hospital"),

        ("h1", "Conditions we treat"),
        ("b", "Diabetic foot ulcers"),
        ("b", "Pressure injuries"),
        ("b", "Surgical wounds"),
        ("b", "Burns and skin tears"),
        ("b", "Leg ulcers and cellulitis"),
        ("b", "Stoma and fistula care"),
        ("b", "Moisture-associated skin damage (MASD)"),

        ("h1", "Our wound management services include"),
        ("b", "Comprehensive wound assessments"),
        ("b", "Advanced wound debridement and dressings"),
        ("b", "EMOLED therapy"),
        ("b", "Patient and caregiver training"),
        ("b", "Dietitian and therapist support for healing"),
        ("b", "Early intervention to prevent amputation or hospital readmission"),
    ],
}

SLH_CHRONIC = {
    "filename": "slh-managing-chronic-wounds.docx",
    "title": "Managing Chronic Wounds",
    "source_title": "St Luke's Hospital — Caring for Chronic Wounds",
    "source_url": "https://www.slh.org.sg/managing-chronic-wounds/",
    "publisher": "St Luke's Hospital, Singapore.",
    "blocks": [
        ("h1", "Caring for chronic wounds"),
        ("p", "At least one in 20 Singaporeans suffer from chronic wounds — a condition where wounds do not "
              "heal properly for more than a month. Examples of chronic wounds include diabetic foot ulcers, "
              "pressure injuries, and arterial and venous leg ulcers."),

        ("h1", "Wounds that do not heal"),
        ("p", "Wounds go through several stages of healing. First, the blood clots to prevent further "
              "bleeding, and a scab forms to protect the wound. As the wound begins to heal over the course "
              "of a few weeks, swelling, pain and discharge will decrease and new tissue will grow over the "
              "wound."),
        ("p", "If the wound does not heal properly after 4 weeks, it will ooze pus and smell, become red "
              "around the wound, become more painful, and may develop into a chronic wound — a condition "
              "that may take years to heal or, at worst case, never heal."),
        ("p", "Chronic wounds can result from excessive pressure on bony parts of the body (pressure "
              "ulcers), damage to the feet or legs due to poor circulation (arterial or venous leg ulcers), "
              "or poor glycaemic control leading to diabetic foot ulcers. Patients with diabetes, "
              "cardiovascular and peripheral vascular disease, and limited mobility are most likely to "
              "develop chronic wounds. Unless addressed with care, these wounds rarely heal on their own."),

        ("h1", "Impact of chronic wounds"),
        ("p", "Patients with these 'hard-to-heal' wounds often experience anxiety and distress from the "
              "prolonged pain and suffering. If not properly treated, they can lead to severe infections, "
              "amputations and even death. Chronic wounds need to be diagnosed and treated early to prevent "
              "such complications."),

        ("h1", "Taking care of wounds"),
        ("p", "Besides treating the underlying cause of the wounds, patients will need proper wound care "
              "and a healthy lifestyle to encourage the wound to heal. Some patients may even need surgery "
              "to remove dead tissue to improve the blood flow and oxygen to the wound."),
        ("p", "Here are three tips to take care of chronic wounds:"),
        ("h2", "1. Keep wounds clean and dry"),
        ("b", "To prevent the wound from being infected, patients and caregivers need to wash their hands "
              "with soap and water before touching the wound and keep a clean dressing on. Dressings absorb "
              "fluid that drains from the wound, keep out germs and protect the wound from further injury."),
        ("b", "When bathing, wounds on the legs should be kept dry with waterproof covers or elevated in a "
              "shower chair to prevent water soaking the wound dressing."),
        ("h2", "2. Protect the wound from trauma or injury"),
        ("p", "Don't let anything touch it or bump against it."),
        ("h2", "3. Check the wound regularly"),
        ("p", "See the doctor immediately if there is bleeding, or if there is more pain or discharge "
              "from the wound."),

        ("h1", "Eating right for healing"),
        ("p", "Eating the right foods gives the wound the building blocks it needs to heal. Patients with "
              "chronic wounds should ensure they get enough protein, vitamin C and zinc in their diets."),
        ("b", "Protein helps build and repair skin and other body tissues. Patients should eat healthy "
              "sources of protein such as lean meat, poultry, fish, eggs, dairy, or soy with each meal."),
        ("b", "Vitamin C plays an important role in helping wounds heal and form new skin tissue. Patients "
              "should eat more fruits and vegetables with good sources of vitamin C such as oranges, kiwis, "
              "berries, tomatoes, and broccoli."),
        ("b", "Zinc is a mineral that is also important for wound healing and formation of new skin tissues. "
              "Foods that are high in zinc include red meat, milk and dairy products, shellfish, beans and "
              "lentils, bread and cereals, and leafy green vegetables."),
        ("b", "Dehydration causes skin to become more fragile and prone to injury. Patients should drink 6 "
              "to 8 glasses of liquids a day unless their doctors advise against it."),
        ("p", "Getting the proper treatment for chronic wounds increases the patient's quality of life and "
              "reduces the burden on caregivers — so don't ignore your wounds."),

        ("h1", "Types of chronic ulcers"),
        ("b", "Arterial ulcers develop as the result of damage to the arteries due to lack of blood flow "
              "to tissue."),
        ("b", "Venous ulcers develop from damage to the veins caused by an insufficient return of blood "
              "back to the heart."),
        ("b", "Peripheral vascular disease (PVD) is a blood circulation disorder that causes the blood "
              "vessels outside of your heart and brain to narrow, block, or spasm."),
    ],
}


def build(doc_spec):
    d = Document()

    # Base font
    normal = d.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title
    d.add_heading(doc_spec["title"], level=0)

    # Source attribution block
    src = d.add_paragraph()
    r = src.add_run(f"Source: {doc_spec['source_title']}")
    r.italic = True
    url = d.add_paragraph()
    ru = url.add_run(doc_spec["source_url"])
    ru.italic = True
    pub = d.add_paragraph()
    rp = pub.add_run(doc_spec["publisher"] + f" Retrieved {RETRIEVED}.")
    rp.italic = True
    rp.font.size = Pt(9)
    rp.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    d.add_paragraph()  # spacer

    for kind, text in doc_spec["blocks"]:
        if kind == "h1":
            d.add_heading(text, level=1)
        elif kind == "h2":
            d.add_heading(text, level=2)
        elif kind == "p":
            d.add_paragraph(text)
        elif kind == "b":
            d.add_paragraph(text, style="List Bullet")
        elif kind == "b2":
            d.add_paragraph(text, style="List Bullet 2")
        elif kind == "quote":
            p = d.add_paragraph()
            rq = p.add_run(text)
            rq.italic = True
        elif kind == "note":
            p = d.add_paragraph()
            rn = p.add_run(text)
            rn.italic = True
            rn.font.size = Pt(9)
            rn.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Footer line reinforcing provenance (helps grounded citation)
    d.add_paragraph()
    foot = d.add_paragraph()
    rf = foot.add_run(f"Source URL: {doc_spec['source_url']}")
    rf.font.size = Pt(9)
    rf.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    out_path = os.path.join(OUT_DIR, doc_spec["filename"])
    d.save(out_path)
    print(f"wrote {out_path}")


for spec in (HEALTHHUB, SLH_TREATMENT, SLH_CHRONIC):
    build(spec)

print("done")
